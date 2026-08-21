"""Document file parsers — extracts text and metadata from various file formats.

Each parser returns a ParsedDocument with the extracted text, per-page content
(when applicable), and file-level metadata.  The FileParser protocol allows
adding new formats without touching existing code.
"""
from __future__ import annotations

import csv
import hashlib
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    """A single page (or logical section) from a document."""

    page_number: int
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Result of parsing a file — full text, per-page breakdown, and metadata."""

    text: str
    pages: list[ParsedPage]
    metadata: dict = field(default_factory=dict)
    file_hash: str = ""


class FileParser(Protocol):
    """Protocol every format-specific parser must satisfy."""

    def can_parse(self, file_type: str) -> bool: ...
    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument: ...


def compute_file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# ── PDF Parser ──────────────────────────────────────────────────────────

class PdfParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "pdf"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("pypdf is required for PDF parsing — pip install pypdf")

        reader = PdfReader(io.BytesIO(file_bytes))
        pages: list[ParsedPage] = []
        all_text_parts: list[str] = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(ParsedPage(page_number=i + 1, text=text))
                all_text_parts.append(text)

        return ParsedDocument(
            text="\n\n".join(all_text_parts),
            pages=pages,
            metadata={"page_count": len(reader.pages), "source_type": "pdf"},
            file_hash=compute_file_hash(file_bytes),
        )


# ── DOCX Parser ─────────────────────────────────────────────────────────

class DocxParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "docx"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx is required for DOCX parsing — pip install python-docx")

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # DOCX doesn't have page numbers; treat as single page
        full_text = "\n\n".join(paragraphs)
        pages = [ParsedPage(page_number=1, text=full_text)] if full_text else []

        return ParsedDocument(
            text=full_text,
            pages=pages,
            metadata={"paragraph_count": len(paragraphs), "source_type": "docx"},
            file_hash=compute_file_hash(file_bytes),
        )


# ── Plain Text Parser ───────────────────────────────────────────────────

class TxtParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "txt"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace").strip()
        pages = [ParsedPage(page_number=1, text=text)] if text else []
        return ParsedDocument(
            text=text,
            pages=pages,
            metadata={"source_type": "txt"},
            file_hash=compute_file_hash(file_bytes),
        )


# ── Markdown Parser ─────────────────────────────────────────────────────

class MarkdownParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "md"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace").strip()

        # Extract headings for metadata
        headings = re.findall(r"^#{1,6}\s+(.+)$", text, re.MULTILINE)

        pages = [ParsedPage(page_number=1, text=text)] if text else []
        return ParsedDocument(
            text=text,
            pages=pages,
            metadata={
                "source_type": "md",
                "headings": headings[:20],  # cap for metadata size
            },
            file_hash=compute_file_hash(file_bytes),
        )


# ── CSV Parser ──────────────────────────────────────────────────────────

class CsvParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "csv"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        # Convert CSV to readable text: header row as column names,
        # each data row as "Column: Value" pairs
        if len(rows) < 2:
            flat = text.strip()
        else:
            headers = rows[0]
            lines: list[str] = []
            for row in rows[1:]:
                pairs = [f"{headers[i]}: {row[i]}" if i < len(headers) else row[i]
                         for i in range(len(row))]
                lines.append(" | ".join(pairs))
            flat = "\n".join(lines)

        pages = [ParsedPage(page_number=1, text=flat)] if flat else []
        return ParsedDocument(
            text=flat,
            pages=pages,
            metadata={"source_type": "csv", "row_count": len(rows) - 1, "column_count": len(rows[0]) if rows else 0},
            file_hash=compute_file_hash(file_bytes),
        )


# ── HTML Parser ─────────────────────────────────────────────────────────

class HtmlParser:
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == "html"

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise RuntimeError("beautifulsoup4 is required for HTML parsing — pip install beautifulsoup4")

        soup = BeautifulSoup(file_bytes, "html.parser")

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        title = soup.title.string if soup.title else None

        pages = [ParsedPage(page_number=1, text=text)] if text else []
        return ParsedDocument(
            text=text,
            pages=pages,
            metadata={"source_type": "html", "title": title},
            file_hash=compute_file_hash(file_bytes),
        )


# ── Script Parser (PowerShell / Bash) ───────────────────────────────────

class ScriptParser:
    """Parses PowerShell (.ps1) and Bash (.sh) scripts, preserving comments
    as documentation context and code blocks as-is."""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() in ("ps1", "sh")

    async def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace").strip()
        file_type = Path(file_name).suffix.lstrip(".")

        # Extract comments as documentation
        if file_type == "ps1":
            comments = re.findall(r"^\s*#(.+)$", text, re.MULTILINE)
            # Also extract block comments <# ... #>
            block_comments = re.findall(r"<#(.*?)#>", text, re.DOTALL)
            comments.extend(block_comments)
        else:  # bash
            comments = re.findall(r"^\s*#(.+)$", text, re.MULTILINE)

        pages = [ParsedPage(page_number=1, text=text)] if text else []
        return ParsedDocument(
            text=text,
            pages=pages,
            metadata={
                "source_type": file_type,
                "language": "powershell" if file_type == "ps1" else "bash",
                "comment_count": len(comments),
            },
            file_hash=compute_file_hash(file_bytes),
        )


# ── Parser Registry ─────────────────────────────────────────────────────

_PARSERS: list[FileParser] = [
    PdfParser(),
    DocxParser(),
    TxtParser(),
    MarkdownParser(),
    CsvParser(),
    HtmlParser(),
    ScriptParser(),
]

SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md",
    ".csv": "csv", ".html": "html", ".htm": "html",
    ".ps1": "ps1", ".sh": "sh",
}


def get_file_type(file_name: str) -> str | None:
    """Resolve a file extension to a normalized type string."""
    ext = Path(file_name).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext)


async def parse_file(file_bytes: bytes, file_name: str, file_type: str) -> ParsedDocument:
    """Route to the appropriate parser and return the extracted content."""
    for parser in _PARSERS:
        if parser.can_parse(file_type):
            return await parser.parse(file_bytes, file_name)
    raise ValueError(f"Unsupported file type: {file_type}")


def validate_file(file_bytes: bytes, file_name: str, max_size_mb: int = 50) -> tuple[bool, str]:
    """Basic file validation: size, extension, and magic bytes checks."""
    # Size check
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"File size ({size_mb:.1f} MB) exceeds maximum ({max_size_mb} MB)"

    # Extension check
    file_type = get_file_type(file_name)
    if file_type is None:
        return False, f"Unsupported file extension: {Path(file_name).suffix}"

    # Basic magic bytes check for common formats
    if file_type == "pdf" and not file_bytes[:5] == b"%PDF-":
        return False, "File content does not match PDF format"

    if file_type == "docx":
        # DOCX files are ZIP archives starting with PK
        if not file_bytes[:2] == b"PK":
            return False, "File content does not match DOCX format"

    return True, ""
